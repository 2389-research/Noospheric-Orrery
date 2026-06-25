import io
import json
import pytest

from src.pipeline.file_extractor import (
    extract_text,
    extract_text_from_notebook,
    ALL_SUPPORTED_EXTENSIONS,
    TEXT_EXTENSIONS,
    IMAGE_EXTENSIONS,
    PDF_EXTENSIONS,
    DOCX_EXTENSIONS,
    NOTEBOOK_EXTENSIONS,
)


# --- extension sets are disjoint where expected ---

def test_extension_sets_consistent():
    assert PDF_EXTENSIONS.isdisjoint(TEXT_EXTENSIONS)
    assert PDF_EXTENSIONS.isdisjoint(IMAGE_EXTENSIONS)
    assert DOCX_EXTENSIONS.isdisjoint(TEXT_EXTENSIONS)
    assert NOTEBOOK_EXTENSIONS.isdisjoint(TEXT_EXTENSIONS)

def test_all_supported_extensions_is_union():
    assert ALL_SUPPORTED_EXTENSIONS == (
        TEXT_EXTENSIONS | IMAGE_EXTENSIONS | PDF_EXTENSIONS | DOCX_EXTENSIONS | NOTEBOOK_EXTENSIONS
    )

def test_doc_not_in_supported_extensions():
    assert ".doc" not in ALL_SUPPORTED_EXTENSIONS
    assert ".doc" not in DOCX_EXTENSIONS


# --- extract_text: plain text types ---

def test_extract_text_txt():
    content = "Hello, world!"
    result = extract_text("file.txt", content.encode("utf-8"))
    assert result == content

def test_extract_text_markdown():
    content = "# Title\n\nSome **bold** text."
    result = extract_text("doc.md", content.encode("utf-8"))
    assert result == content

def test_extract_text_python():
    content = "def foo():\n    return 42\n"
    result = extract_text("script.py", content.encode("utf-8"))
    assert result == content

def test_extract_text_invalid_utf8_replaced():
    bad_bytes = b"good text \xbf\xfe bad bytes"
    result = extract_text("file.txt", bad_bytes)
    assert "good text" in result
    assert isinstance(result, str)

def test_extract_text_unsupported_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text("archive.zip", b"data")

def test_extract_text_no_extension_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text("noextension", b"data")

def test_extract_text_doc_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text("legacy.doc", b"data")


# --- extract_text_from_notebook ---

def _make_notebook(cells: list) -> bytes:
    nb = {"nbformat": 4, "nbformat_minor": 5, "metadata": {}, "cells": cells}
    return json.dumps(nb).encode("utf-8")

def test_notebook_markdown_cell():
    nb = _make_notebook([{"cell_type": "markdown", "source": ["# Hello\n", "World"], "metadata": {}}])
    result = extract_text_from_notebook(nb)
    assert "# Hello\nWorld" in result

def test_notebook_code_cell():
    nb = _make_notebook([{"cell_type": "code", "source": ["x = 1\n", "print(x)"], "metadata": {}, "outputs": []}])
    result = extract_text_from_notebook(nb)
    assert "```python" in result
    assert "x = 1" in result

def test_notebook_code_cell_with_stream_output():
    nb = _make_notebook([{
        "cell_type": "code",
        "source": ["print('hi')"],
        "metadata": {},
        "outputs": [{"output_type": "stream", "text": ["hi\n"]}],
    }])
    result = extract_text_from_notebook(nb)
    assert "Output:" in result
    assert "hi" in result

def test_notebook_code_cell_with_display_output():
    nb = _make_notebook([{
        "cell_type": "code",
        "source": ["42"],
        "metadata": {},
        "outputs": [{"output_type": "execute_result", "data": {"text/plain": ["42"]}, "metadata": {}}],
    }])
    result = extract_text_from_notebook(nb)
    assert "42" in result

def test_notebook_empty_cells_skipped():
    nb = _make_notebook([
        {"cell_type": "markdown", "source": ["   "], "metadata": {}},
        {"cell_type": "code", "source": [""], "metadata": {}, "outputs": []},
        {"cell_type": "markdown", "source": ["real content"], "metadata": {}},
    ])
    result = extract_text_from_notebook(nb)
    assert result.strip() == "real content"

def test_notebook_mixed_cells_ordering():
    nb = _make_notebook([
        {"cell_type": "markdown", "source": ["intro"], "metadata": {}},
        {"cell_type": "code", "source": ["x = 1"], "metadata": {}, "outputs": []},
        {"cell_type": "markdown", "source": ["conclusion"], "metadata": {}},
    ])
    result = extract_text_from_notebook(nb)
    parts = result.split("\n\n")
    assert parts[0] == "intro"
    assert "x = 1" in parts[1]
    assert parts[-1] == "conclusion"

def test_notebook_empty():
    nb = _make_notebook([])
    result = extract_text_from_notebook(nb)
    assert result == ""

def test_extract_text_routes_notebook():
    nb = _make_notebook([{"cell_type": "markdown", "source": ["hello"], "metadata": {}}])
    result = extract_text("analysis.ipynb", nb)
    assert "hello" in result


# --- PDF and DOCX: smoke tests with real minimal files ---

def test_extract_text_pdf_smoke():
    pytest.importorskip("pypdf")
    reportlab = pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas as rl_canvas
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(100, 750, "Hello PDF")
    c.save()
    result = extract_text("doc.pdf", buf.getvalue())
    assert "Hello" in result

def test_extract_text_pdf_blank_raises():
    pypdf = pytest.importorskip("pypdf")
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    buf = io.BytesIO()
    writer.write(buf)
    with pytest.raises(ValueError, match="no extractable text"):
        extract_text("blank.pdf", buf.getvalue())

def test_extract_text_pdf_corrupt_raises():
    pytest.importorskip("pypdf")
    with pytest.raises(Exception):
        extract_text("corrupt.pdf", b"%PDF-1.4 \x00garbage\xff\xfe")

def test_extract_text_docx_smoke():
    docx = pytest.importorskip("docx")
    import docx as docx_lib
    doc = docx_lib.Document()
    doc.add_paragraph("Test paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text("test.docx", buf.getvalue())
    assert "Test paragraph" in result

def test_extract_text_docx_empty_paragraphs_skipped():
    docx = pytest.importorskip("docx")
    import docx as docx_lib
    doc = docx_lib.Document()
    doc.add_paragraph("")
    doc.add_paragraph("content")
    doc.add_paragraph("   ")
    buf = io.BytesIO()
    doc.save(buf)
    result = extract_text("test.docx", buf.getvalue())
    assert result == "content"


def test_extract_text_docx_zipbomb_raises():
    pytest.importorskip("docx")
    import zipfile
    from src.pipeline.file_extractor import _DOCX_UNZIP_LIMIT
    # Build a zip whose actual uncompressed content exceeds the limit.
    # ZIP_STORED so file_size in the header matches the real payload size.
    chunk = b"A" * (1024 * 1024)  # 1 MB per member
    n_members = (_DOCX_UNZIP_LIMIT // len(chunk)) + 1
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_STORED) as zf:
        for i in range(n_members):
            zf.writestr(f"chunk{i}.bin", chunk)
    with pytest.raises(ValueError, match="exceeds limit"):
        extract_text("bomb.docx", buf.getvalue())
